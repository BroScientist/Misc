from shutil import copyfile
from docx import Document
import datetime

# visit parent file, increment numbers, save and duplicate (reminder: order is date, number, weekday)
parent_path = '/Users/apple/Documents/Diaries/mother_of_all_diaries.docx'
document = Document(parent_path)

today = datetime.datetime.today()
date_str = f'{today.month}月{today.day}日'

# check if diary file is already up-to-date (avoid crontab updating the file multiple times)
if document.paragraphs[0].text == date_str:
    print('file already up to date')

else:
    # extract diary number and increment it by one
    text = document.paragraphs[1].text
    digits = [c for c in text if c.isdigit()]
    diary_num = int(''.join(digits)) + 1
    diary_num_str = f'日记#{diary_num}'

    # get chinese weekday based on current weekday
    chinese_weekdays = ['星期一', '星期二', '星期三', '星期四', '星期五', '星期六', '星期日']
    weekday_str = chinese_weekdays[today.weekday()]

    # modify the file with the updated strings
    document.paragraphs[0].text = date_str
    document.paragraphs[1].text = diary_num_str
    document.paragraphs[2].text = weekday_str

    # save the parent file for future use
    document.save(parent_path)

    # save the new file to the main diary folder with appropriate diary name (date + ' ' + year)
    diary_name_str = f'{date_str} {today.year}'
    diary_folder = '/Users/apple/Documents/Diaries/我的日记2020'
    new_diary_path = f'{diary_folder}/{diary_name_str}.docx'
    copyfile(parent_path, new_diary_path)
    print('successfully generated diary template')

# schedule to launch daily (may require multiple launches and checks to ensure files gets generated)


